"""
Document processing service for extracting text from various file formats.
Supports PDF, DOCX, plain text, and images (with OCR).
"""
from typing import Dict, Optional
import io
import os
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

# Document processing libraries
try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    from PIL import Image
except ImportError:
    Image = None

# PaddleOCR for handwriting recognition
try:
    print("\n" + "="*80)
    print("Attempting to import PaddleOCR service...")
    from app.services.paddleocr_service import get_paddleocr_service
    PADDLEOCR_AVAILABLE = True
    print("PADDLEOCR SERVICE IMPORTED SUCCESSFULLY")
    print("="*80 + "\n")
    logger.info("PaddleOCR service imported successfully")
except ImportError as e:
    PADDLEOCR_AVAILABLE = False
    print("\n" + "="*80)
    print(f"PADDLEOCR SERVICE NOT AVAILABLE: {e}")
    print("="*80 + "\n")
    logger.warning(f"PaddleOCR service not available: {e}")


class DocumentProcessor:
    """Service for processing and extracting text from documents."""

    SUPPORTED_FORMATS = {
        'pdf': ['.pdf'],
        'docx': ['.docx', '.doc'],
        'text': ['.txt'],
        'image': ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']
    }

    def __init__(self, llm_orchestrator=None):
        """Initialize the document processor."""
        print("\n" + "="*80)
        print("INITIALIZING DOCUMENT PROCESSOR")
        print("="*80)
        self._check_dependencies()
        self._llm_orchestrator = llm_orchestrator

        # Initialize OCR strategy
        self.ocr_strategy = os.getenv("OCR_STRATEGY", "auto").lower()
        print(f"OCR Strategy: {self.ocr_strategy}")
        print(f"PaddleOCR Available: {PADDLEOCR_AVAILABLE}")
        print("="*80 + "\n")
        logger.info(f"DocumentProcessor initialized - OCR strategy: {self.ocr_strategy}")
        logger.info(f"PaddleOCR available: {PADDLEOCR_AVAILABLE}")

    def _check_dependencies(self):
        """Check if required libraries are available."""
        if pdfplumber is None:
            logger.warning("pdfplumber not installed. PDF processing will fail.")
        if DocxDocument is None:
            logger.warning("python-docx not installed. DOCX processing will fail.")
        if Image is None:
            logger.warning("PIL not installed. Image OCR will fail.")

    def _get_llm_orchestrator(self):
        """Get LLM orchestrator instance (lazy import to avoid circular deps)."""
        if self._llm_orchestrator is None:
            from app.services.llm_orchestrator import get_llm_orchestrator
            self._llm_orchestrator = get_llm_orchestrator()
        return self._llm_orchestrator

    def get_file_type(self, filename: str) -> Optional[str]:
        """
        Determine file type from filename extension.

        Args:
            filename: Name of the file

        Returns:
            File type (pdf, docx, text, image) or None if unsupported
        """
        ext = Path(filename).suffix.lower()
        for file_type, extensions in self.SUPPORTED_FORMATS.items():
            if ext in extensions:
                return file_type
        return None

    async def extract_text(
        self,
        file_data: bytes,
        filename: str,
        file_type: Optional[str] = None
    ) -> Dict[str, any]:
        """
        Extract text from a document.

        Args:
            file_data: File content as bytes
            filename: Original filename
            file_type: Optional file type override

        Returns:
            dict with extracted_text, page_count, metadata
        """
        try:
            # Determine file type
            if file_type is None:
                file_type = self.get_file_type(filename)

            if file_type is None:
                raise ValueError(f"Unsupported file format: {filename}")

            print(f"\nextract_text called - File: {filename}, Type: {file_type}")

            # Extract text based on file type
            if file_type == 'pdf':
                print("Routing to _extract_from_pdf...")
                return await self._extract_from_pdf(file_data)
            elif file_type == 'docx':
                print("Routing to _extract_from_docx...")
                return await self._extract_from_docx(file_data)
            elif file_type == 'text':
                print("Routing to _extract_from_text...")
                return await self._extract_from_text(file_data)
            elif file_type == 'image':
                print("Routing to _extract_from_image...")
                return await self._extract_from_image(file_data)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")

        except Exception as e:
            logger.error(f"Error extracting text from {filename}: {e}")
            raise

    async def _extract_from_pdf(self, file_data: bytes) -> Dict[str, any]:
        """Extract text from PDF using pdfplumber, with OCR fallback for scanned PDFs."""
        print("\n_extract_from_pdf called")
        if pdfplumber is None:
            raise RuntimeError("pdfplumber is not installed")

        try:
            text_content = []
            page_count = 0

            with pdfplumber.open(io.BytesIO(file_data)) as pdf:
                page_count = len(pdf.pages)

                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(page_text)

            extracted_text = "\n\n".join(text_content)
            print(f"pdfplumber extracted {len(extracted_text)} chars")
            print(f"First 200 chars: {extracted_text[:200] if extracted_text else 'EMPTY'}")

            # If PaddleOCR is enabled, always use OCR for better handwriting recognition
            # Even if pdfplumber extracts text, it may be gibberish from handwritten content
            if self.ocr_strategy in ["auto", "paddleocr"]:
                print(f"OCR strategy is '{self.ocr_strategy}' - forcing OCR for handwriting")
                return await self._extract_from_pdf_with_ocr(file_data)

            # If no text was extracted, the PDF might be scanned (image-based)
            # Fall back to OCR
            if len(extracted_text.strip()) < 50:
                print(f"Less than 50 chars - triggering OCR fallback")
                logger.warning(f"PDF appears to be scanned (extracted only {len(extracted_text)} chars). Attempting OCR...")
                return await self._extract_from_pdf_with_ocr(file_data)

            print(f"Returning pdfplumber text (no OCR)")
            return {
                "extracted_text": extracted_text,
                "page_count": page_count,
                "char_count": len(extracted_text),
                "word_count": len(extracted_text.split()),
                "format": "pdf",
                "ocr_used": False,
                "page_images": []
            }

        except Exception as e:
            logger.error(f"Error extracting PDF: {e}")
            raise

    async def _extract_from_pdf_with_ocr(self, file_data: bytes) -> Dict[str, any]:
        """Extract text from scanned PDF using best available OCR method."""
        if Image is None:
            raise RuntimeError(
                "PIL is required for PDF to image conversion. "
                "Install with: pip install Pillow"
            )

        try:
            import pdf2image
        except ImportError:
            raise RuntimeError(
                "pdf2image is required for PDF OCR. "
                "Install with: pip install pdf2image"
            )

        try:
            logger.info("Converting PDF to images for OCR processing...")

            # Convert PDF pages to images
            images = pdf2image.convert_from_bytes(file_data)

            # Determine OCR method
            use_paddleocr = self._should_use_paddleocr()

            text_content = []
            page_images: list = []
            for i, image in enumerate(images):
                logger.info(f"Processing page {i + 1}/{len(images)}...")

                # Convert to RGB if necessary
                if image.mode != 'RGB':
                    image = image.convert('RGB')

                # Convert PIL Image to bytes
                img_byte_arr = io.BytesIO()
                image.save(img_byte_arr, format='PNG')
                img_bytes = img_byte_arr.getvalue()

                # Retain page image for vision evaluation
                page_images.append(img_bytes)

                # Perform OCR with three-layer fallback
                try:
                    logger.info(f"Processing page {i + 1}/{len(images)} with three-layer OCR fallback...")
                    page_text, method = await self._extract_with_three_layer_fallback(img_bytes)

                    if page_text.strip():
                        text_content.append(page_text)
                        logger.info(f"Page {i + 1}: Extracted {len(page_text)} chars using {method}")
                    else:
                        logger.warning(f"Page {i + 1}: No text extracted")
                except Exception as e:
                    logger.error(f"Page {i + 1}: All OCR layers failed - {e}")
                    continue

            extracted_text = "\n\n".join(text_content)

            logger.info(
                f"OCR completed. Extracted {len(extracted_text)} characters "
                f"from {len(images)} pages"
            )

            return {
                "extracted_text": extracted_text,
                "page_count": len(images),
                "char_count": len(extracted_text),
                "word_count": len(extracted_text.split()),
                "format": "pdf",
                "ocr_used": True,
                "page_images": page_images,
                "metadata": {
                    "pages_processed": len(images),
                    "method": "PaddleOCR (Handwriting)" if use_paddleocr else "GLM-OCR via Ollama",
                    "strategy": self.ocr_strategy
                }
            }

        except Exception as e:
            logger.error(f"Error extracting PDF with OCR: {e}")
            raise

    async def _extract_from_docx(self, file_data: bytes) -> Dict[str, any]:
        """Extract text from DOCX using python-docx."""
        if DocxDocument is None:
            raise RuntimeError("python-docx is not installed")

        try:
            doc = DocxDocument(io.BytesIO(file_data))

            # Extract text from paragraphs
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text for cell in row.cells)
                    if row_text.strip():
                        text_content.append(row_text)

            extracted_text = "\n".join(text_content).replace('\xa0', ' ')

            return {
                "extracted_text": extracted_text,
                "page_count": len(doc.sections) if hasattr(doc, 'sections') else 1,
                "char_count": len(extracted_text),
                "word_count": len(extracted_text.split()),
                "format": "docx",
                "page_images": []
            }

        except Exception as e:
            logger.error(f"Error extracting DOCX: {e}")
            raise

    async def _extract_from_text(self, file_data: bytes) -> Dict[str, any]:
        """Extract text from plain text file."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            extracted_text = None

            for encoding in encodings:
                try:
                    extracted_text = file_data.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue

            if extracted_text is None:
                raise ValueError("Could not decode text file with any supported encoding")

            return {
                "extracted_text": extracted_text.replace('\xa0', ' '),
                "page_count": 1,
                "char_count": len(extracted_text),
                "word_count": len(extracted_text.split()),
                "format": "text",
                "page_images": []
            }

        except Exception as e:
            logger.error(f"Error extracting text: {e}")
            raise

    async def _extract_from_image(self, file_data: bytes) -> Dict[str, any]:
        """Extract text from image using best available OCR method."""
        if Image is None:
            raise RuntimeError("PIL is not installed. Install with: pip install Pillow")

        try:
            # Open image to get metadata
            image = Image.open(io.BytesIO(file_data))
            image_size = image.size
            image_mode = image.mode

            # Convert to RGB if necessary for consistency
            if image.mode != 'RGB':
                image = image.convert('RGB')

                # Convert back to bytes
                img_byte_arr = io.BytesIO()
                image.save(img_byte_arr, format='PNG')
                file_data = img_byte_arr.getvalue()

            # Perform OCR with three-layer fallback
            extracted_text, method = await self._extract_with_three_layer_fallback(file_data)

            return {
                "extracted_text": extracted_text,
                "page_count": 1,
                "char_count": len(extracted_text),
                "word_count": len(extracted_text.split()),
                "format": "image",
                "ocr_used": True,
                "page_images": [file_data],
                "metadata": {
                    "image_size": image_size,
                    "image_mode": image_mode,
                    "method": method,
                    "strategy": self.ocr_strategy
                }
            }

        except Exception as e:
            logger.error(f"Error extracting text from image: {e}")
            raise

    async def validate_file(
        self,
        file_data: bytes,
        filename: str,
        max_size: Optional[int] = None
    ) -> Dict[str, any]:
        """
        Validate a file before processing.

        Args:
            file_data: File content as bytes
            filename: Original filename
            max_size: Maximum allowed file size in bytes

        Returns:
            dict with is_valid, file_type, file_size, errors
        """
        errors = []
        file_size = len(file_data)

        # Check file type
        file_type = self.get_file_type(filename)
        if file_type is None:
            errors.append(f"Unsupported file format: {Path(filename).suffix}")

        # Check file size
        if max_size and file_size > max_size:
            errors.append(f"File size ({file_size} bytes) exceeds maximum ({max_size} bytes)")

        # Check if file is empty
        if file_size == 0:
            errors.append("File is empty")

        return {
            "is_valid": len(errors) == 0,
            "file_type": file_type,
            "file_size": file_size,
            "errors": errors
        }

    def _should_use_paddleocr(self) -> bool:
        """Determine if PaddleOCR should be used based on strategy."""
        print(f"\nChecking PaddleOCR - Strategy: {self.ocr_strategy}, Available: {PADDLEOCR_AVAILABLE}")
        logger.info(f"Checking PaddleOCR availability - Strategy: {self.ocr_strategy}, Available: {PADDLEOCR_AVAILABLE}")

        if self.ocr_strategy == "ollama":
            print("OCR strategy is 'ollama' - skipping PaddleOCR")
            logger.info("OCR strategy is 'ollama' - skipping PaddleOCR")
            return False

        if self.ocr_strategy == "paddleocr":
            print("OCR strategy is 'paddleocr' - using PaddleOCR")
            logger.info("OCR strategy is 'paddleocr' - using PaddleOCR")
            return True

        # Auto mode: use PaddleOCR if available
        if self.ocr_strategy == "auto":
            print("Auto mode - checking PaddleOCR availability...")
            if PADDLEOCR_AVAILABLE:
                print("Getting PaddleOCR service...")
                paddleocr = get_paddleocr_service()
                is_available = paddleocr.is_available()
                print(f"{'Available' if is_available else 'Not Available'} PaddleOCR service.is_available() = {is_available}")
                logger.info(f"PaddleOCR service availability: {is_available}")
                return is_available
            else:
                print("PaddleOCR not available - falling back to Ollama")
                logger.warning("PaddleOCR not available - falling back to Ollama")
            return False

        print("Unknown strategy - not using PaddleOCR")
        return False

    async def _extract_with_three_layer_fallback(self, image_data: bytes) -> tuple[str, str]:
        """
        Extract text using three-layer OCR fallback strategy.

        Returns:
            tuple: (extracted_text, method_used)
        """
        # Layer 1: Try PaddleOCR library first (best quality, fastest)
        if self.ocr_strategy in ["auto", "paddleocr"]:
            if PADDLEOCR_AVAILABLE:
                paddleocr = get_paddleocr_service()
                if paddleocr.is_available():
                    try:
                        logger.info("Layer 1: Trying PaddleOCR library...")
                        result = await paddleocr.extract_text_from_image(image_data)
                        text = result.get("extracted_text", "")
                        if text.strip():
                            logger.info(f"Layer 1: PaddleOCR library extracted {len(text)} chars")
                            return text, "PaddleOCR (Library)"
                        else:
                            logger.warning("Layer 1: PaddleOCR extracted no text")
                    except Exception as e:
                        logger.warning(f"Layer 1 failed (PaddleOCR library): {e}")

        # Layer 2: Try PaddleOCR-VL via Ollama (OCR-focused fallback)
        if self.ocr_strategy == "auto":
            try:
                logger.info("Layer 2: Trying PaddleOCR-VL via Ollama...")
                llm = self._get_llm_orchestrator()
                text = await llm.extract_text_with_paddleocr_vl(image_data)
                if text.strip():
                    logger.info(f"Layer 2: PaddleOCR-VL extracted {len(text)} chars")
                    return text, "PaddleOCR-VL (Ollama)"
                else:
                    logger.warning("Layer 2: PaddleOCR-VL extracted no text")
            except Exception as e:
                logger.warning(f"Layer 2 failed (PaddleOCR-VL via Ollama): {e}")

        # Layer 3: GLM-OCR via Ollama (last resort)
        try:
            logger.info("Layer 3: Trying GLM-OCR via Ollama (last resort)...")
            from app.config import settings
            llm = self._get_llm_orchestrator()
            text = await llm.extract_text_from_image_with_ocr(
                image_data,
                model=settings.OLLAMA_FALLBACK_MODEL
            )
            logger.info(f"Layer 3: GLM-OCR extracted {len(text)} chars")
            return text, "GLM-OCR (Ollama - Fallback)"
        except Exception as e:
            logger.error(f"All 3 OCR layers failed! Last error: {e}")
            raise RuntimeError("All OCR methods failed. Unable to extract text from image.")

    async def _extract_with_paddleocr(self, image_data: bytes) -> str:
        """Extract text using PaddleOCR library (Layer 1 only)."""
        paddleocr = get_paddleocr_service()
        result = await paddleocr.extract_text_from_image(image_data)
        return result.get("extracted_text", "")

    async def _extract_with_paddleocr_vl(self, image_data: bytes) -> str:
        """Extract text using PaddleOCR-VL via Ollama (Layer 2 only)."""
        llm = self._get_llm_orchestrator()
        return await llm.extract_text_with_paddleocr_vl(image_data)

    async def _extract_with_glm_ocr(self, image_data: bytes) -> str:
        """Extract text using GLM-OCR via Ollama (Layer 3 only)."""
        from app.config import settings
        llm = self._get_llm_orchestrator()
        return await llm.extract_text_from_image_with_ocr(
            image_data,
            model=settings.OLLAMA_FALLBACK_MODEL
        )


# Singleton instance
_document_processor: Optional[DocumentProcessor] = None


def get_document_processor() -> DocumentProcessor:
    """Get the document processor singleton."""
    global _document_processor
    if _document_processor is None:
        _document_processor = DocumentProcessor()
    return _document_processor
