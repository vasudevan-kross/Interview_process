"""
LLM-based Resume Parser and Matcher.

Uses Ollama to parse resumes and match with job descriptions.
No PyTorch required - uses pdfplumber for PDF text extraction and LLM for parsing.
No OCR dependencies needed for resume matching.
"""

import logging
import io
import json
from typing import Dict, Any, List, Optional
from pathlib import Path
from pydantic import BaseModel, Field
from app.services.llm_orchestrator import LLMOrchestrator

logger = logging.getLogger(__name__)

# Document processing libraries (pdfplumber only, no OCR)
try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None


# ── Pydantic schemas for LLM grammar-constrained output ───────────────────

class _ExperienceEntry(BaseModel):
    title: Optional[str] = None
    company: Optional[str] = None
    duration: Optional[str] = None
    description: Optional[str] = None

class _EducationEntry(BaseModel):
    degree: Optional[str] = None
    institution: Optional[str] = None
    year: Optional[str] = None

class _ParsedResume(BaseModel):
    name: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    summary: Optional[str] = None
    skills: List[str] = []
    experience: List[_ExperienceEntry] = []
    education: List[_EducationEntry] = []
    years_of_experience: Optional[int] = None

class _MatchResult(BaseModel):
    match_score: int = Field(0, ge=0, le=100)
    overall_assessment: Optional[str] = None
    strengths: List[str] = []
    weaknesses: List[str] = []
    missing_skills: List[str] = []
    matching_skills: List[str] = []
    experience_match: Optional[str] = None
    education_match: Optional[str] = None
    recommendation: Optional[str] = None
    reasoning: Optional[str] = None


class ResumeParserLLM:
    """Parse resumes and match with job descriptions using LLM."""

    def __init__(self):
        self.llm = LLMOrchestrator()

    def _extract_text(self, file_data: bytes, filename: str) -> str:
        """
        Extract text from a resume file using pdfplumber (PDF),
        python-docx (DOCX), or plain decode (TXT).

        No OCR is used — resumes are expected to be text-based documents.

        Args:
            file_data: File content as bytes
            filename: Original filename

        Returns:
            Extracted text string
        """
        ext = Path(filename).suffix.lower()

        if ext == '.pdf':
            return self._extract_from_pdf(file_data)
        elif ext in ('.docx', '.doc'):
            return self._extract_from_docx(file_data)
        elif ext == '.txt':
            return self._extract_from_text(file_data)
        else:
            raise ValueError(f"Unsupported resume format: {ext}. Supported: .pdf, .docx, .doc, .txt")

    def _extract_from_pdf(self, file_data: bytes) -> str:
        """Extract text from PDF using pdfplumber only (no OCR)."""
        if pdfplumber is None:
            raise RuntimeError("pdfplumber is not installed. Install with: pip install pdfplumber")

        text_content = []
        with pdfplumber.open(io.BytesIO(file_data)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_content.append(page_text)

        return "\n\n".join(text_content)

    def _extract_from_docx(self, file_data: bytes) -> str:
        """Extract text from DOCX using python-docx."""
        if DocxDocument is None:
            raise RuntimeError("python-docx is not installed. Install with: pip install python-docx")

        doc = DocxDocument(io.BytesIO(file_data))
        text_content = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)

        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text for cell in row.cells)
                if row_text.strip():
                    text_content.append(row_text)

        return "\n".join(text_content)

    def _extract_from_text(self, file_data: bytes) -> str:
        """Extract text from plain text file."""
        for encoding in ('utf-8', 'latin-1', 'cp1252'):
            try:
                return file_data.decode(encoding)
            except UnicodeDecodeError:
                continue
        raise ValueError("Could not decode text file with any supported encoding")

    async def parse_resume(self, file_data: bytes, filename: str) -> Dict[str, Any]:
        """
        Extract structured information from resume.

        Args:
            file_data: Resume file content
            filename: Resume filename

        Returns:
            dict with name, email, phone, skills, experience, education, summary
        """
        try:
            # Extract text from resume using pdfplumber (no OCR)
            resume_text = self._extract_text(file_data, filename)

            if not resume_text or len(resume_text.strip()) < 50:
                raise ValueError("Could not extract text from resume")

            logger.info(f"Extracted {len(resume_text)} chars from resume: {filename}")

            # Parse resume using LLM with schema-constrained output + retry
            prompt = f"""Parse the following resume and extract structured information.

EXAMPLE INPUT:
John Smith | john@example.com | +1-555-0100
Software Engineer at Acme Corp (2021-2023): Built REST APIs using Python/FastAPI.
B.Sc Computer Science, MIT, 2021. Skills: Python, FastAPI, PostgreSQL, Docker.

EXAMPLE OUTPUT:
{{"name":"John Smith","email":"john@example.com","phone":"+1-555-0100","summary":"Software engineer with 2 years of backend experience.","skills":["Python","FastAPI","PostgreSQL","Docker"],"experience":[{{"title":"Software Engineer","company":"Acme Corp","duration":"2021-2023","description":"Built REST APIs using Python/FastAPI."}}],"education":[{{"degree":"B.Sc Computer Science","institution":"MIT","year":"2021"}}],"years_of_experience":2}}

Now parse this resume:
{resume_text[:6000]}

Return ONLY the JSON object. Use null for missing fields, empty array [] for missing lists."""

            parsed_data = await self.llm.generate_with_retry(
                prompt=prompt,
                schema_class=_ParsedResume,
                schema=_ParsedResume.model_json_schema(),
                temperature=0.1,
                max_tokens=1024,
            )

            logger.info(f"Successfully parsed resume: {parsed_data.get('name', 'Unknown')}")

            return {
                'parsed_data': parsed_data,
                'raw_text': resume_text,
                'filename': filename
            }

        except Exception as e:
            logger.error(f"Error parsing resume {filename}: {e}")
            raise

    async def match_with_job(
        self,
        resume_data: Dict[str, Any],
        job_description: str,
        job_title: str
    ) -> Dict[str, Any]:
        """
        Match resume with job description and generate score.

        Args:
            resume_data: Parsed resume data from parse_resume()
            job_description: Job description text
            job_title: Job title

        Returns:
            dict with match_score (0-100), strengths, weaknesses, recommendation
        """
        try:
            parsed = resume_data.get('parsed_data') or {}
            candidate_name = parsed.get('name', 'Candidate')

            prompt = f"""Analyze how well this candidate matches the job requirements.

Job Title: {job_title}

Job Description:
{job_description[:2000]}

Candidate Information:
- Name: {parsed.get('name', 'N/A')}
- Experience: {parsed.get('years_of_experience', 0)} years
- Skills: {', '.join(parsed.get('skills', [])[:20])}
- Education: {json.dumps(parsed.get('education', [])[:3])}
- Recent Experience: {json.dumps(parsed.get('experience', [])[:3])}

Return a JSON object with these fields:
- match_score: integer 0-100 (0=no match, 100=perfect match)
- overall_assessment: "Strong match", "Good match", "Partial match", or "Weak match"
- strengths: array of strings (specific matching points)
- weaknesses: array of strings (gaps)
- missing_skills: array of skill name strings
- matching_skills: array of skill name strings
- experience_match: "Exceeds requirements", "Meets requirements", or "Below requirements"
- education_match: "Strong match", "Good match", "Acceptable", or "Not relevant"
- recommendation: "Strong recommend", "Recommend", "Consider", or "Not recommended"
- reasoning: 2-3 sentence explanation

Return ONLY the JSON object."""

            match_result = await self.llm.generate_with_retry(
                prompt=prompt,
                schema_class=_MatchResult,
                schema=_MatchResult.model_json_schema(),
                temperature=0.1,
                max_tokens=1024,
            )

            # Clamp score to valid range regardless of schema enforcement
            match_result['match_score'] = min(100, max(0, match_result.get('match_score', 0)))

            logger.info(f"Match score for {candidate_name}: {match_result['match_score']}/100")

            return match_result

        except Exception as e:
            logger.error(f"Error matching resume with job: {e}")
            raise


# Singleton instance
_resume_parser_llm = None


def get_resume_parser_llm() -> ResumeParserLLM:
    """Get singleton instance of ResumeParserLLM."""
    global _resume_parser_llm
    if _resume_parser_llm is None:
        _resume_parser_llm = ResumeParserLLM()
    return _resume_parser_llm
