"""
Coding Interviews API endpoints.

Provides REST API for:
- Creating and managing interviews
- AI question generation
- Candidate submission tracking
- Code auto-save and evaluation
- Anti-cheating monitoring
"""

import logging
import json
import re
import zipfile
from datetime import datetime
from io import BytesIO, StringIO
from fastapi import APIRouter, BackgroundTasks, HTTPException, status, Request, Body, Depends, UploadFile, File, Form
from fastapi.responses import StreamingResponse
from typing import List, Optional

from app.schemas.coding_interviews import (
    InterviewCreate,
    GenerateQuestionsRequest,
    StartSubmissionRequest,
    SaveCodeRequest,
    SubmitInterviewRequest,
    TrackActivityRequest,
    BulkImportResponse,
    InterviewCandidateResponse,
    CandidateDecisionUpdate,
    CandidateUpdate,
    EvaluatorNotesUpdate,
    InterviewUpdate,
    BulkDecisionRequest,
    BulkDeleteCandidatesRequest,
)
from app.services.coding_interview_service import get_coding_interview_service
from app.services.question_generator import get_question_generator, DOMAIN_REGISTRY
from app.services.document_processor import get_document_processor
from app.services.resume_auto_processor import get_resume_auto_processor
from app.services.llm_orchestrator import get_llm_orchestrator
from app.services.voice_interview_service import get_voice_interview_service
from app.db.supabase_client import get_supabase
from app.auth.dependencies import get_current_user_id, get_current_org_context, OrgContext
from app.auth.permissions import require_permission
from pydantic import BaseModel

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/coding-interviews", tags=["coding-interviews"])


@router.get("/domains", summary="List available interview domains")
async def list_domains():
    """
    Return the domain registry config so the frontend can render domain/tool dropdowns dynamically.
    Adding a new domain requires only a change to DOMAIN_REGISTRY in question_generator.py.
    """
    return {"domains": DOMAIN_REGISTRY}


@router.post("", summary="Create coding interview")
async def create_interview(
    request: InterviewCreate,
    ctx: OrgContext = Depends(require_permission('interview:create'))
):
    """
    Create a new coding or testing interview with questions.

    Returns:
        - interview_id
        - access_token (for shareable link)
        - shareable_link
        - questions
    """
    try:
        service = get_coding_interview_service()

        # Convert questions to dict
        questions_data = [q.dict() for q in request.questions]

        result = await service.create_interview(
            title=request.title,
            description=request.description or "",
            scheduled_start_time=request.scheduled_start_time,
            scheduled_end_time=request.scheduled_end_time,
            programming_language=request.programming_language,
            interview_type=request.interview_type,
            questions_data=questions_data,
            user_id=ctx.user_id,
            grace_period_minutes=request.grace_period_minutes,
            resume_required=request.resume_required,
            allowed_languages=request.allowed_languages,
            bond_terms=request.bond_terms,
            bond_document_url=request.bond_document_url,
            require_signature=request.require_signature,
            bond_years=request.bond_years,
            bond_timing=request.bond_timing,
            job_id=request.job_id,
            org_id=ctx.org_id
        )

        return result

    except Exception as e:
        logger.error(f"Error creating interview: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to create interview: {str(e)}"
        )


@router.post("/generate-questions", summary="Generate questions with AI")
async def generate_questions(
    request: GenerateQuestionsRequest,
    ctx: OrgContext = Depends(require_permission('interview:create'))
):
    """
    Generate interview questions using AI (Ollama codellama:7b).

    Supports:
    - Coding questions (algorithms, development)
    - Testing questions (test cases, automation code)
    - Auto-detects testing roles from job description
    """
    try:
        generator = get_question_generator()

        # --- Auto-detect domain from job description keywords ---
        jd_lower = request.job_description.lower()

        KEYWORD_DOMAINS = [
            ('devops', ['devops', 'docker', 'kubernetes', 'k8s', 'terraform', 'ansible', 'ci/cd', 'jenkins', 'pipeline', 'infrastructure', 'cloud engineer', 'site reliability', 'sre']),
            ('sql', ['sql', 'database', 'dba', 'postgres', 'mysql', 'oracle', 'schema design', 'query optimiz']),
            ('data_science', ['data scientist', 'machine learning', 'ml engineer', 'data analyst', 'pandas', 'numpy', 'sklearn', 'deep learning', 'nlp', 'computer vision']),
            ('testing', ['tester', 'testing', 'qa ', 'quality assurance', 'sdet', 'test engineer', 'test analyst', 'automation tester', 'manual tester', 'selenium', 'playwright', 'cypress', 'appium', 'jmeter', 'regression', 'smoke test']),
        ]

        effective_type = request.interview_type
        for domain, keywords in KEYWORD_DOMAINS:
            if any(kw in jd_lower for kw in keywords):
                if effective_type == 'coding':  # Only auto-switch if still on default
                    effective_type = domain
                    logger.info(f"Auto-detected '{domain}' role from job description")
                break

        if effective_type == 'both':
            # Legacy: mix coding + testing
            coding_count = request.num_questions // 2
            testing_count = request.num_questions - coding_count
            coding_qs = await generator.generate_coding_questions(
                job_description=request.job_description,
                difficulty=request.difficulty,
                num_questions=coding_count,
                programming_language=request.programming_language or 'python'
            )
            testing_qs = await generator.generate_testing_questions(
                job_description=request.job_description,
                difficulty=request.difficulty,
                num_questions=testing_count,
                test_framework=request.test_framework or 'manual-test-cases'
            )
            questions = coding_qs + testing_qs
        else:
            # Use registry dispatch for all other domains
            questions = await generator.generate_questions_for_domain(
                domain=effective_type,
                job_description=request.job_description,
                difficulty=request.difficulty,
                num_questions=request.num_questions,
                domain_tool=request.domain_tool,
                programming_language=request.programming_language,
                test_framework=request.test_framework,
            )

        return {
            'questions': questions,
            'count': len(questions),
            'detected_type': effective_type,
        }

    except Exception as e:
        logger.error(f"Error generating questions: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to generate questions: {str(e)}"
        )


# ---------------------------------------------------------------------------
# Voice Interview Creation
# ---------------------------------------------------------------------------

class VoiceSessionRequest(BaseModel):
    message: str
    session_state: dict = {}
    user_timezone_offset: int = 330  # Minutes ahead of UTC (default: IST +5:30)


@router.get("/voice-session/start", summary="Get opening message for voice interview creation")
async def voice_session_start(
    ctx: OrgContext = Depends(require_permission('interview:create'))
):
    """
    Returns the opening greeting message and initial session state for voice-driven
    interview creation. Call this once when the user opens the voice modal.
    """
    service = get_voice_interview_service()
    return service.get_opening_message()


@router.post("/voice-session", summary="Process one turn in voice interview creation")
async def voice_session(
    request: VoiceSessionRequest,
    ctx: OrgContext = Depends(require_permission('interview:create'))
):
    """
    Process a single voice turn. The caller sends the latest speech transcript
    and the full session_state from the previous response. Returns the agent's
    reply, updated session_state, and (when done=true) the created interview details.
    """
    try:
        service = get_voice_interview_service()
        return await service.process_turn(
            message=request.message,
            session_state=request.session_state,
            user_id=ctx.user_id,
            user_timezone_offset=request.user_timezone_offset,
        )
    except Exception as e:
        logger.error(f"Voice session error: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Voice session failed: {str(e)}"
        )

@router.post("/extract-questions", summary="Extract questions from document")
async def extract_questions_from_document(
    file: UploadFile = File(...),
    programming_language: str = Form('python'),
    interview_type: str = Form('coding'),
    difficulty: str = Form('medium'),
    ctx: OrgContext = Depends(require_permission('interview:create'))
):
    """
    Extract interview questions from uploaded document.

    Supports: PDF, Word (DOCX), Images (JPG, PNG), Excel (XLSX), CSV

    Process:
    1. Extract text using OCR/document processing
    2. Parse questions using LLM
    3. Return structured questions for review
    """
    try:
        # Validate file type
        allowed_extensions = ['pdf', 'docx', 'doc', 'jpg', 'jpeg', 'png', 'xlsx', 'csv']
        file_ext = file.filename.split('.')[-1].lower()

        if file_ext not in allowed_extensions:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Unsupported file type: {file_ext}. Allowed: {', '.join(allowed_extensions)}"
            )

        logger.info(f"Extracting questions from {file.filename} ({file_ext})")

        # Read file content
        content = await file.read()

        # Extract text from document
        doc_processor = get_document_processor()
        extraction_result = await doc_processor.extract_text(
            file_data=content,
            filename=file.filename,
            file_type=file_ext
        )

        extracted_text = extraction_result.get('extracted_text', '')

        if not extracted_text or len(extracted_text.strip()) < 10:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Could not extract text from document. Please ensure document is readable."
            )

        logger.info(f"Extracted {len(extracted_text)} characters from document ({extraction_result.get('page_count', 0)} pages)")

        # Parse questions using LLM
        llm = get_llm_orchestrator()

        system_prompt = """You are an expert at extracting interview questions from documents and structuring them as JSON.
Always return valid JSON arrays only, without any additional text or explanation."""

        prompt = f"""Extract coding interview questions from the following document text and return them as a JSON array.

Document Text:
{extracted_text[:4000]}

Return ONLY a JSON array with this exact structure (no markdown, no explanations):
[
  {{
    "question_text": "Full question text here",
    "difficulty": "{difficulty}",
    "marks": 10,
    "topics": ["relevant", "topics"],
    "starter_code": "",
    "solution_code": ""
  }}
]

Rules:
- Extract ALL questions from the document
- If marks not specified, use: easy=5-10, medium=10-20, hard=20-30
- Keep exact question wording from document
- Return valid JSON array only"""

        result = await llm.generate_completion(
            prompt=prompt,
            system_prompt=system_prompt,
            temperature=0.3  # Low temperature for more consistent JSON output
        )

        # Parse LLM response as JSON
        try:
            # Extract JSON from response (handle markdown code blocks and various formats)
            json_text = result['response'].strip()

            # Remove markdown code blocks
            if '```json' in json_text:
                json_text = json_text.split('```json')[1].split('```')[0].strip()
            elif '```' in json_text:
                json_text = json_text.split('```')[1].split('```')[0].strip()

            # Try to find JSON array in the response
            if '[' in json_text:
                start = json_text.index('[')
                end = json_text.rindex(']') + 1
                json_text = json_text[start:end]

            json_text = json_text.strip()

            if not json_text:
                raise ValueError("Empty response from LLM")

            questions = json.loads(json_text)

            if not isinstance(questions, list):
                raise ValueError("Expected array of questions")

            logger.info(f"Successfully extracted {len(questions)} questions")

            return {
                'questions': questions,
                'count': len(questions),
                'extracted_text_length': len(extracted_text)
            }

        except json.JSONDecodeError as e:
            logger.error(f"Failed to parse LLM response as JSON: {e}")
            logger.error(f"LLM Response: {result.get('response', '')[:500]}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to parse questions from document. LLM response was not valid JSON. Please try again or check the document format."
            )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error extracting questions from document: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to extract questions: {str(e)}"
        )


@router.get("", summary="List interviews")
async def list_interviews(
    ctx: OrgContext = Depends(require_permission('interview:view')),
    status_filter: Optional[str] = None,
    limit: int = 50,
    offset: int = 0
):
    """List all interviews for the current organization."""
    try:
        client = get_supabase()
        from datetime import datetime, timezone

        # Auto-update expired interviews before listing
        # Find interviews that are still 'scheduled' or 'in_progress' but past their end time
        active_result = client.table('coding_interviews').select(
            'id, status, scheduled_end_time, grace_period_minutes'
        ).eq('org_id', ctx.org_id).in_(
            'status', ['scheduled', 'in_progress']
        ).execute()

        now = datetime.now(timezone.utc)
        if active_result.data:
            for interview in active_result.data:
                end_time_str = interview.get('scheduled_end_time')
                if not end_time_str:
                    continue

                # Parse end time
                try:
                    end_time = datetime.fromisoformat(end_time_str.replace('Z', '+00:00'))
                    # Ensure timezone-aware (treat naive as UTC)
                    if end_time.tzinfo is None:
                        end_time = end_time.replace(tzinfo=timezone.utc)
                except (ValueError, AttributeError):
                    continue

                grace_minutes = interview.get('grace_period_minutes', 0) or 0
                from datetime import timedelta
                effective_end = end_time + timedelta(minutes=grace_minutes)

                if now > effective_end:
                    # Time has expired — update status
                    new_status = 'expired' if interview['status'] == 'scheduled' else 'completed'
                    try:
                        client.table('coding_interviews').update({
                            'status': new_status
                        }).eq('id', interview['id']).execute()
                        logger.info(f"Auto-updated interview {interview['id']} to {new_status}")
                    except Exception as update_err:
                        logger.warning(f"Failed to auto-update interview status: {update_err}")

        # Now fetch the list with updated statuses
        query = client.table('coding_interviews').select('*').eq('org_id', ctx.org_id).is_('deleted_at', 'null').order('created_at', desc=True).limit(limit).offset(offset)

        if status_filter:
            query = query.eq('status', status_filter)

        result = query.execute()

        interviews = result.data or []
        
        # Get submission counts
        if interviews:
            interview_ids = [i['id'] for i in interviews]
            subs_result = client.table('coding_submissions').select('interview_id').in_('interview_id', interview_ids).in_('status', ['submitted', 'auto_submitted', 'evaluated']).execute()
            counts = {}
            if subs_result.data:
                for sub in subs_result.data:
                    counts[sub['interview_id']] = counts.get(sub['interview_id'], 0) + 1
                    
            # Calculate duration_minutes and attach submission_count dynamically for each interview
            for interview in interviews:
                interview['submission_count'] = counts.get(interview['id'], 0)
                try:
                    start_time = datetime.fromisoformat(interview['scheduled_start_time'].replace('Z', '+00:00'))
                    end_time = datetime.fromisoformat(interview['scheduled_end_time'].replace('Z', '+00:00'))
                    interview['duration_minutes'] = int((end_time - start_time).total_seconds() / 60)
                except (ValueError, KeyError, TypeError):
                    interview['duration_minutes'] = None

        return {
            'interviews': interviews,
            'count': len(interviews)
        }

    except Exception as e:
        logger.error(f"Error listing interviews: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to list interviews: {str(e)}"
        )


@router.get("/{interview_id}", summary="Get interview details")
async def get_interview(
    interview_id: str,
    ctx: OrgContext = Depends(require_permission('interview:view'))
):
    """Get interview details with questions."""
    try:
        client = get_supabase()

        # Get interview and verify org membership
        interview_result = client.table('coding_interviews').select('*').eq(
            'id', interview_id
        ).eq('org_id', ctx.org_id).execute()

        if not interview_result.data or len(interview_result.data) == 0:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Interview not found"
            )

        interview = interview_result.data[0]

        # Auto-update status if end time has passed
        from datetime import datetime, timezone, timedelta
        if interview.get('status') in ('scheduled', 'in_progress'):
            end_time_str = interview.get('scheduled_end_time')
            if end_time_str:
                try:
                    end_time = datetime.fromisoformat(end_time_str.replace('Z', '+00:00'))
                    # Ensure timezone-aware (treat naive as UTC)
                    if end_time.tzinfo is None:
                        end_time = end_time.replace(tzinfo=timezone.utc)
                    grace_minutes = interview.get('grace_period_minutes', 0) or 0
                    effective_end = end_time + timedelta(minutes=grace_minutes)
                    now = datetime.now(timezone.utc)

                    if now > effective_end:
                        new_status = 'expired' if interview['status'] == 'scheduled' else 'completed'
                        client.table('coding_interviews').update({
                            'status': new_status
                        }).eq('id', interview_id).execute()
                        interview['status'] = new_status
                        logger.info(f"Auto-updated interview {interview_id} to {new_status}")
                except (ValueError, AttributeError):
                    pass

        # Get questions
        questions_result = client.table('coding_questions').select('*').eq(
            'interview_id', interview_id
        ).order('question_number').execute()

        interview['questions'] = questions_result.data or []

        # Calculate duration_minutes dynamically from scheduled times
        try:
            start_time = datetime.fromisoformat(interview['scheduled_start_time'].replace('Z', '+00:00'))
            end_time = datetime.fromisoformat(interview['scheduled_end_time'].replace('Z', '+00:00'))
            interview['duration_minutes'] = int((end_time - start_time).total_seconds() / 60)
        except (ValueError, KeyError, TypeError):
            interview['duration_minutes'] = None

        return interview

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting interview: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to get interview: {str(e)}"
        )


@router.delete("/{interview_id}", summary="Delete interview")
async def delete_interview(
    interview_id: str,
    ctx: OrgContext = Depends(require_permission('interview:create'))
):
    """Delete an interview (cascade deletes questions, submissions)."""
    try:
        client = get_supabase()

        # Soft delete
        result = client.table('coding_interviews').update(
            {"deleted_at": datetime.utcnow().isoformat()}
        ).eq('id', interview_id).eq('org_id', ctx.org_id).is_('deleted_at', 'null').execute()

        if not result.data or len(result.data) == 0:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Interview not found"
            )

        return {'message': 'Interview deleted successfully'}

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error deleting interview: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to delete interview: {str(e)}"
        )


@router.get("/submissions/{submission_id}/risk-score", summary="Get submission risk score")
async def get_submission_risk_score(
    submission_id: str,
    ctx: OrgContext = Depends(require_permission('interview:view'))
):
    """
    Get comprehensive risk score for a submission based on anti-cheating activities.

    Returns:
    - total_risk_score: Sum of all risk points
    - risk_level: low/medium/high/critical
    - activity_counts: Count of each activity type
    - high_risk_activities: List of concerning activities
    - flagged_events: Specific flagged events
    """
    try:
        service = get_coding_interview_service()
        client = get_supabase()

        # Verify ownership
        submission_result = client.table('coding_submissions').select(
            'id, interview_id'
        ).eq('id', submission_id).execute()

        if not submission_result.data:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Submission not found"
            )

        submission = submission_result.data[0]

        # Check if user's org owns the interview
        interview_result = client.table('coding_interviews').select('id').eq(
            'id', submission['interview_id']
        ).eq('org_id', ctx.org_id).execute()

        if not interview_result.data:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Not authorized to view this submission"
            )

        # Calculate risk score
        risk_score = await service.get_submission_risk_score(submission_id)

        return risk_score

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting risk score: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to get risk score: {str(e)}"
        )


@router.get("/join/{access_token}", summary="Join interview (public)")
async def join_interview(access_token: str):
    """
    Public endpoint for candidates to join interview.
    Validates token and returns interview details (without solutions).
    """
    try:
        service = get_coding_interview_service()

        interview = await service.get_interview_by_token(access_token)

        return interview

    except ValueError as e:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=str(e)
        )
    except Exception as e:
        logger.error(f"Error joining interview: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to join interview"
        )


@router.post("/start", summary="Start submission (public)")
async def start_submission(
    request: StartSubmissionRequest,
    interview_id: str,
    http_request: Request
):
    """
    Public endpoint for candidates to start interview.
    Creates submission record and validates time window.
    """
    try:
        service = get_coding_interview_service()

        # Get IP and user agent
        ip_address = http_request.client.host
        user_agent = http_request.headers.get('user-agent', '')

        result = await service.start_submission(
            interview_id=interview_id,
            candidate_name=request.candidate_name,
            candidate_email=request.candidate_email,
            candidate_phone=request.candidate_phone,
            ip_address=ip_address,
            user_agent=user_agent,
            preferred_language=request.preferred_language
        )

        return result

    except ValueError as e:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=str(e)
        )
    except Exception as e:
        logger.error(f"Error starting submission: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to start submission"
        )


@router.post("/save-code", summary="Auto-save code (public)")
async def save_code(request: SaveCodeRequest):
    """
    Public endpoint for auto-saving code.
    Called every 30 seconds from frontend.
    """
    try:
        service = get_coding_interview_service()

        result = await service.save_code(
            submission_id=request.submission_id,
            question_id=request.question_id,
            code=request.code,
            programming_language=request.programming_language
        )

        return result

    except Exception as e:
        logger.error(f"Error saving code: {type(e).__name__}: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to save code: {type(e).__name__}: {str(e)}"
        )


@router.post("/submit", summary="Submit interview (public)")
async def submit_interview(
    request: SubmitInterviewRequest,
    req: Request,
    background_tasks: BackgroundTasks
):
    """
    Public endpoint for submitting interview.
    Finalizes submission and triggers evaluation.
    Accepts optional signature data and terms acceptance.

    Also triggers background resume processing to link uploaded resumes to pipeline.
    """
    try:
        service = get_coding_interview_service()

        # Get IP address for audit trail
        client_ip = req.client.host if req.client else None

        result = await service.submit_interview(
            submission_id=request.submission_id,
            auto_submit=False,
            signature_data=request.signature_data,
            terms_accepted=request.terms_accepted,
            client_ip=client_ip
        )

        # Add background task to process resume (if uploaded)
        background_tasks.add_task(
            process_resume_background,
            request.submission_id
        )

        return result

    except ValueError as e:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=str(e)
        )
    except Exception as e:
        logger.error(f"Error submitting interview: {type(e).__name__}: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to submit interview: {type(e).__name__}: {str(e)}"
        )


async def process_resume_background(submission_id: str):
    """Background task to process resume after submission."""
    try:
        processor = get_resume_auto_processor()
        await processor.process_coding_interview_resume(submission_id)
    except Exception as e:
        logger.error(f"Background resume processing failed for submission {submission_id}: {e}")
        # Don't raise - background task failure shouldn't affect submission


@router.post("/activity", summary="Track activity (public)")
async def track_activity(request: TrackActivityRequest):
    """
    Public endpoint for tracking anti-cheating events.
    Logs tab switches, copy/paste, etc.
    """
    try:
        service = get_coding_interview_service()

        await service.track_activity(
            submission_id=request.submission_id,
            activity_type=request.activity_type,
            question_id=request.question_id,
            metadata=request.metadata
        )

        return {'status': 'logged'}

    except Exception as e:
        logger.error(f"Error tracking activity: {e}")
        # Don't fail the request if activity tracking fails
        return {'status': 'failed', 'error': str(e)}


@router.post("/upload-resume", summary="Upload resume (public)")
async def upload_resume(
    submission_id: str = Form(...),
    file: UploadFile = File(...)
):
    """
    Public endpoint for candidates to upload their resume.
    Stores file in Supabase Storage and updates submission record.
    """
    try:
        # Validate file type
        allowed_extensions = ['pdf', 'docx', 'doc']
        file_ext = file.filename.split('.')[-1].lower() if file.filename else ''

        if file_ext not in allowed_extensions:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Unsupported file type: {file_ext}. Allowed: {', '.join(allowed_extensions)}"
            )

        # Read file content
        content = await file.read()

        if len(content) > 10 * 1024 * 1024:  # 10MB limit
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="File too large. Maximum size is 10MB."
            )

        client = get_supabase()

        # Verify submission exists
        submission_result = client.table('coding_submissions').select('id, interview_id').eq(
            'id', submission_id
        ).execute()

        if not submission_result.data:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Submission not found"
            )

        # Upload to Supabase Storage
        import uuid
        storage_path = f"coding-resumes/{submission_id}/{uuid.uuid4()}.{file_ext}"

        try:
            client.storage.from_('documents').upload(
                storage_path,
                content,
                file_options={"content-type": file.content_type or "application/octet-stream"}
            )
        except Exception as storage_err:
            # If bucket doesn't exist, try creating it
            logger.warning(f"Storage upload failed, trying to create bucket: {storage_err}")
            try:
                client.storage.create_bucket('documents', options={"public": False})
                client.storage.from_('documents').upload(
                    storage_path,
                    content,
                    file_options={"content-type": file.content_type or "application/octet-stream"}
                )
            except Exception as retry_err:
                logger.error(f"Storage upload retry failed: {retry_err}")
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail="Failed to upload resume to storage"
                )

        # Update submission with resume path
        from datetime import datetime
        client.table('coding_submissions').update({
            'resume_path': storage_path,
            'resume_uploaded_at': datetime.now().isoformat()
        }).eq('id', submission_id).execute()

        logger.info(f"Resume uploaded for submission {submission_id}: {storage_path}")

        return {
            'status': 'uploaded',
            'resume_path': storage_path,
            'filename': file.filename
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error uploading resume: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to upload resume: {str(e)}"
        )


@router.get("/submissions/{submission_id}/resume", summary="Download resume")
async def get_submission_resume(
    submission_id: str,
    ctx: OrgContext = Depends(require_permission('interview:view'))
):
    """Get a signed URL for viewing the candidate's uploaded resume."""
    try:
        client = get_supabase()

        # Get submission
        submission_result = client.table('coding_submissions').select(
            'resume_path, interview_id'
        ).eq('id', submission_id).execute()

        if not submission_result.data:
            raise HTTPException(status_code=404, detail="Submission not found")

        submission = submission_result.data[0]

        # Verify interview org membership
        interview_result = client.table('coding_interviews').select('id').eq(
            'id', submission['interview_id']
        ).eq('org_id', ctx.org_id).execute()

        if not interview_result.data:
            raise HTTPException(status_code=403, detail="Not authorized")

        resume_path = submission.get('resume_path')
        if not resume_path:
            raise HTTPException(status_code=404, detail="No resume uploaded for this submission")

        # Generate signed URL (valid for 1 hour)
        signed_url = client.storage.from_('documents').create_signed_url(
            resume_path, 3600
        )

        return {
            'resume_url': signed_url.get('signedURL') or signed_url.get('signedUrl', ''),
            'resume_path': resume_path
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting resume: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to get resume: {str(e)}"
        )


@router.get("/{interview_id}/submissions", summary="List submissions")
async def list_submissions(
    interview_id: str,
    ctx: OrgContext = Depends(require_permission('interview:view'))
):
    """List all submissions for an interview."""
    try:
        client = get_supabase()

        # Verify interview org membership
        interview_result = client.table('coding_interviews').select('id').eq(
            'id', interview_id
        ).eq('org_id', ctx.org_id).execute()

        if not interview_result.data:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Interview not found"
            )

        # Get submissions
        submissions_result = client.table('coding_submissions').select('*').eq(
            'interview_id', interview_id
        ).order('submitted_at', desc=True).execute()

        return {
            'submissions': submissions_result.data or [],
            'count': len(submissions_result.data) if submissions_result.data else 0
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error listing submissions: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to list submissions"
        )


@router.get("/submissions/{submission_id}", summary="Get submission details")
async def get_submission(
    submission_id: str,
    ctx: OrgContext = Depends(require_permission('interview:view'))
):
    """Get detailed submission with all answers and evaluations."""
    try:
        client = get_supabase()

        # Get submission
        submission_result = client.table('coding_submissions').select('*').eq(
            'id', submission_id
        ).execute()

        if not submission_result.data:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Submission not found"
            )

        submission = submission_result.data[0]

        # Verify interview org membership and fetch details like bond_terms
        interview_result = client.table('coding_interviews').select('id, title, bond_terms').eq(
            'id', submission['interview_id']
        ).eq('org_id', ctx.org_id).execute()

        if not interview_result.data:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Not authorized to view this submission"
            )
            
        submission['interview'] = interview_result.data[0]

        # Fetch questions to attach question_text to each answer - order by question_number
        questions_result = client.table('coding_questions').select(
            'id, question_number, question_text, difficulty, marks, topics'
        ).eq('interview_id', submission['interview_id']).order('question_number').execute()

        questions_list = questions_result.data or []
        questions_map = {q['id']: q for q in questions_list}

        # Get answers with evaluations
        answers_result = client.table('coding_answers').select('*').eq(
            'submission_id', submission_id
        ).execute()

        answers_data = answers_result.data or []
        answers_dict = {a['question_id']: a for a in answers_data}

        # Build ordered answers list based on question_number
        ordered_answers = []
        for q in questions_list:
            answer = answers_dict.get(q['id'])
            if answer:
                answer['question_text'] = q.get('question_text', '')
                answer['question_marks'] = q.get('marks', 0)
                answer['question_difficulty'] = q.get('difficulty', 'medium')
                answer['question_topics'] = q.get('topics', [])
                answer['question_number'] = q.get('question_number')
                ordered_answers.append(answer)

        submission['answers'] = ordered_answers

        # Get activity log
        activities_result = client.table('session_activities').select('*').eq(
            'submission_id', submission_id
        ).order('timestamp').execute()

        submission['activities'] = activities_result.data or []

        return submission

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting submission: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to get submission"
        )


@router.post("/submissions/{submission_id}/evaluate", summary="Re-evaluate submission")
async def reevaluate_submission(
    submission_id: str,
    ctx: OrgContext = Depends(require_permission('interview:evaluate'))
):
    """Manually trigger re-evaluation of a submission."""
    try:
        service = get_coding_interview_service()

        # Verify org membership
        client = get_supabase()
        submission_result = client.table('coding_submissions').select('interview_id').eq(
            'id', submission_id
        ).execute()

        if not submission_result.data:
            raise HTTPException(status_code=404, detail="Submission not found")

        interview_result = client.table('coding_interviews').select('id').eq(
            'id', submission_result.data[0]['interview_id']
        ).eq('org_id', ctx.org_id).execute()

        if not interview_result.data:
            raise HTTPException(status_code=403, detail="Not authorized")

        # Re-evaluate
        result = await service.evaluate_submission(submission_id)

        return result

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error re-evaluating submission: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to re-evaluate submission"
        )


async def _run_bulk_evaluation(interview_id: str, submissions: list) -> None:
    """Background task: evaluate all submissions without holding the HTTP connection."""
    service = get_coding_interview_service()
    total = len(submissions)
    evaluated = 0
    failed = 0
    logger.info(f"[BG] Starting bulk evaluation for {total} submissions in interview {interview_id}")
    for submission in submissions:
        submission_id = submission['id']
        try:
            result = await service.evaluate_submission(submission_id)
            evaluated += 1
            logger.info(f"[BG] ✅ Evaluated {submission_id}: {result.get('percentage')}%")
        except Exception as e:
            failed += 1
            logger.error(f"[BG] ❌ Failed to evaluate {submission_id}: {e}")
    logger.info(f"[BG] Bulk evaluation done: {evaluated} ok, {failed} failed out of {total}")


@router.post("/{interview_id}/evaluate-all", summary="Bulk evaluate all submissions")
async def evaluate_all_submissions(
    interview_id: str,
    background_tasks: BackgroundTasks,
    ctx: OrgContext = Depends(require_permission('interview:evaluate'))
):
    """Start bulk evaluation in background and return immediately to avoid connection timeout."""
    try:
        client = get_supabase()

        # Verify org membership
        interview_result = client.table('coding_interviews').select('id').eq(
            'id', interview_id
        ).eq('org_id', ctx.org_id).execute()

        if not interview_result.data:
            raise HTTPException(status_code=403, detail="Not authorized to access this interview")

        # Get all submitted submissions for this interview
        submissions_result = client.table('coding_submissions').select('id, status').eq(
            'interview_id', interview_id
        ).in_('status', ['submitted', 'auto_submitted', 'evaluated']).execute()

        if not submissions_result.data:
            return {
                "message": "No submissions to evaluate",
                "total": 0,
                "evaluated": 0,
                "failed": 0,
                "results": [],
                "status": "done"
            }

        submissions = submissions_result.data
        total = len(submissions)
        logger.info(f"Queuing bulk evaluation for {total} submissions in interview {interview_id}")

        background_tasks.add_task(_run_bulk_evaluation, interview_id, submissions)

        return {
            "message": f"Evaluation started for {total} submissions. Results will be available shortly.",
            "total": total,
            "evaluated": 0,
            "failed": 0,
            "results": [],
            "status": "processing"
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error starting bulk evaluation: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to start bulk evaluation"
        )


@router.post(
    "/{interview_id}/candidates/bulk",
    response_model=BulkImportResponse,
    summary="Bulk import candidates from Excel/CSV"
)
async def bulk_import_candidates(
    interview_id: str,
    file: UploadFile = File(...),
    ctx: OrgContext = Depends(require_permission('interview:create'))
):
    """
    Upload an Excel (.xlsx / .xls) or CSV file to pre-register candidates for an interview.
    Required column: name. Optional: email, phone. Column names are case-insensitive.
    """
    try:
        import pandas as pd
        from io import BytesIO

        filename = file.filename or ''
        if filename.endswith('.csv'):
            content = await file.read()
            df = pd.read_csv(BytesIO(content))
        elif filename.endswith(('.xlsx', '.xls')):
            content = await file.read()
            df = pd.read_excel(BytesIO(content))
        else:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Only CSV and Excel files (.csv, .xlsx, .xls) are supported"
            )

        # Normalize column names
        df.columns = [c.strip().lower() for c in df.columns]

        if 'name' not in df.columns:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="File must contain a 'name' column"
            )

        # Column alias resolution (case-insensitive, already lowercased above)
        cols = set(df.columns)
        email_col = next((c for c in [
            'email', 'email id', 'emailid', 'email address', 'e-mail', 'mail'
        ] if c in cols), None)
        phone_col = next((c for c in [
            'phone', 'mobile', 'mobile no', 'mobile number', 'phone number',
            'contact', 'contact number', 'cell', 'telephone'
        ] if c in cols), None)

        import math

        # Build candidate list
        candidates = []
        for _, row in df.iterrows():
            def safe_val(v):
                """Convert pandas NaN / empty string to None."""
                if v is None:
                    return None
                try:
                    if math.isnan(float(v)):
                        return None
                except (TypeError, ValueError):
                    pass
                s = str(v).strip()
                return s if s else None

            candidates.append({
                'name': str(row.get('name', '') or '').strip(),
                'email': safe_val(row[email_col] if email_col and email_col in row else None),
                'phone': safe_val(row[phone_col] if phone_col and phone_col in row else None),
            })

        service = get_coding_interview_service()
        result = await service.bulk_import_candidates(
            interview_id=interview_id,
            candidates=candidates,
            user_id=ctx.user_id
        )
        return result

    except HTTPException:
        raise
    except ValueError as e:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail=str(e))
    except Exception as e:
        logger.error(f"Error importing candidates: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to import candidates: {str(e)}"
        )


@router.get(
    "/{interview_id}/candidates",
    summary="List all candidates for an interview"
)
async def get_interview_candidates(
    interview_id: str,
    ctx: OrgContext = Depends(require_permission('interview:view'))
):
    """
    Return unified candidate list: pre-registered + walk-in submissions.
    Requires interview org membership.
    """
    try:
        client = get_supabase()

        # Verify org membership
        interview_result = client.table('coding_interviews').select('id, title, access_token, total_marks').eq(
            'id', interview_id
        ).eq('org_id', ctx.org_id).execute()

        if not interview_result.data:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Interview not found"
            )

        service = get_coding_interview_service()
        candidates = await service.get_interview_candidates(interview_id)
        interview = interview_result.data[0]

        return {
            'interview_id': interview_id,
            'interview_title': interview.get('title'),
            'access_token': interview.get('access_token'),
            'interview_total_marks': interview.get('total_marks'),
            'candidates': candidates,
            'total': len(candidates),
            'submitted': sum(1 for c in candidates if c['submitted']),
            'advanced': sum(1 for c in candidates if c['decision'] == 'advanced'),
            'rejected': sum(1 for c in candidates if c['decision'] == 'rejected'),
            'hold': sum(1 for c in candidates if c['decision'] == 'hold'),
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting candidates: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to get candidates: {str(e)}"
        )


@router.patch(
    "/submissions/{submission_id}/decision",
    summary="Set decision on a submission"
)
async def set_submission_decision(
    submission_id: str,
    body: CandidateDecisionUpdate,
    ctx: OrgContext = Depends(require_permission('interview:evaluate'))
):
    """
    Mark a candidate submission as advanced / rejected / hold / pending.
    Requires org membership of the parent interview.
    """
    try:
        client = get_supabase()

        # Verify org membership via interview
        sub_result = client.table('coding_submissions').select('interview_id').eq(
            'id', submission_id
        ).execute()

        if not sub_result.data:
            raise HTTPException(status_code=404, detail="Submission not found")

        interview_result = client.table('coding_interviews').select('id').eq(
            'id', sub_result.data[0]['interview_id']
        ).eq('org_id', ctx.org_id).execute()

        if not interview_result.data:
            raise HTTPException(status_code=403, detail="Not authorized")

        allowed_decisions = {'advanced', 'rejected', 'hold', 'pending'}
        if body.decision not in allowed_decisions:
            raise HTTPException(
                status_code=400,
                detail=f"Invalid decision. Must be one of: {', '.join(allowed_decisions)}"
            )

        service = get_coding_interview_service()
        result = await service.set_submission_decision(
            submission_id=submission_id,
            decision=body.decision,
            notes=body.notes,
            decided_by=ctx.user_id
        )
        return result

    except HTTPException:
        raise
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))
    except Exception as e:
        logger.error(f"Error setting decision: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to set decision: {str(e)}"
        )


@router.patch(
    "/{interview_id}/candidates/{candidate_id}",
    summary="Edit a pre-registered candidate"
)
async def update_candidate(
    interview_id: str,
    candidate_id: str,
    body: CandidateUpdate,
    ctx: OrgContext = Depends(require_permission('interview:create'))
):
    """Update name, email, phone of a pre-registered candidate."""
    try:
        service = get_coding_interview_service()
        result = await service.update_interview_candidate(
            interview_id=interview_id,
            candidate_id=candidate_id,
            name=body.name,
            email=body.email,
            phone=body.phone,
            user_id=ctx.user_id
        )
        return result
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))
    except Exception as e:
        logger.error(f"Error updating candidate: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to update candidate: {str(e)}")


@router.delete(
    "/{interview_id}/candidates/{candidate_id}",
    summary="Remove a pre-registered candidate"
)
async def delete_candidate(
    interview_id: str,
    candidate_id: str,
    ctx: OrgContext = Depends(require_permission('interview:create'))
):
    """Delete a pre-registered candidate from the interview list."""
    try:
        service = get_coding_interview_service()
        await service.delete_interview_candidate(
            interview_id=interview_id,
            candidate_id=candidate_id,
            user_id=ctx.user_id
        )
        return {'message': 'Candidate removed successfully'}
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))
    except Exception as e:
        logger.error(f"Error deleting candidate: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to delete candidate: {str(e)}")


def _sanitize_name(name: str) -> str:
    """Replace unsafe characters with underscores for ZIP entry names."""
    sanitized = re.sub(r'[^\w\s\-]', '_', name)
    sanitized = re.sub(r'\s+', ' ', sanitized).strip()
    return sanitized or 'Unknown'


@router.get("/{interview_id}/export-csv", summary="Export submissions as CSV")
async def export_submissions_csv(
    interview_id: str,
    ctx: OrgContext = Depends(require_permission('interview:view'))
):
    """Export candidate submissions and question scores as a CSV file."""
    try:
        service = get_coding_interview_service()
        
        # Get interview title for filename
        interview = await service.get_interview(interview_id)
        title = _sanitize_name(interview.get('title', 'Assessment'))
        filename = f"{title}_submissions.csv"
        
        csv_content = await service.export_submissions_csv(
            interview_id=interview_id,
            user_id=ctx.user_id,
            org_id=ctx.org_id
        )
        
        return StreamingResponse(
            StringIO(csv_content),
            media_type="text/csv",
            headers={
                "Content-Disposition": f"attachment; filename={filename}"
            }
        )
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))
    except Exception as e:
        logger.error(f"Error exporting CSV: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to export CSV: {str(e)}")


@router.get("/{interview_id}/export", summary="Export submissions as ZIP")
async def export_submissions_zip(
    interview_id: str,
    ctx: OrgContext = Depends(require_permission('interview:view'))
):
    """
    Download all submitted candidates' resumes + answers as a ZIP archive.

    ZIP structure:
        {assessment_title}/
            {candidate_name}/
                resume.{ext}   (if uploaded)
                answers.docx   (styled Word document)
    """
    try:
        client = get_supabase()

        # Verify interview org membership
        interview_result = client.table('coding_interviews').select(
            'id, title, org_id'
        ).eq('id', interview_id).eq('org_id', ctx.org_id).execute()

        if not interview_result.data:
            raise HTTPException(status_code=404, detail="Interview not found")

        interview = interview_result.data[0]

        title = interview.get('title', 'Assessment')
        safe_title = _sanitize_name(title)

        # Fetch questions for this interview - order by question_number
        questions_result = client.table('coding_questions').select(
            'id, question_number, question_text, difficulty, marks'
        ).eq('interview_id', interview_id).order('question_number').execute()
        
        questions_list = (questions_result.data or [])
        questions = {q['id']: q for q in questions_list}

        # Fetch submitted submissions
        subs_result = client.table('coding_submissions').select(
            'id, candidate_name, candidate_email, total_marks_obtained, percentage, '
            'submitted_at, status, resume_path, candidate_decision'
        ).eq('interview_id', interview_id).in_(
            'status', ['submitted', 'auto_submitted', 'evaluated']
        ).execute()

        all_submissions = subs_result.data or []
        
        # Keep only evaluated ones
        submissions = [s for s in all_submissions if s.get('total_marks_obtained') is not None]

        # Pre-process to identify duplicate candidate names
        name_counts = {}
        for sub in submissions:
            candidate_name = sub.get('candidate_name') or 'Unknown'
            safe_base = _sanitize_name(candidate_name).lower()
            name_counts[safe_base] = name_counts.get(safe_base, 0) + 1

        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        DECISION_COLORS = {
            'advanced': RGBColor(0x16, 0xa3, 0x4a),   # green
            'rejected': RGBColor(0xdc, 0x26, 0x26),   # red
            'hold':     RGBColor(0xd9, 0x77, 0x06),   # amber
            'pending':  RGBColor(0x6b, 0x72, 0x80),   # gray
        }

        def _add_label_value(para, label: str, value: str):
            """Bold label, normal value on the same paragraph."""
            run = para.add_run(f"{label}: ")
            run.bold = True
            run.font.size = Pt(11)
            val_run = para.add_run(value)
            val_run.font.size = Pt(11)

        def _build_docx(sub: dict, answers: list) -> bytes:
            doc = Document()

            # ── Page margins ────────────────────────────────────────────────
            for section in doc.sections:
                section.top_margin    = Inches(0.9)
                section.bottom_margin = Inches(0.9)
                section.left_margin   = Inches(1.0)
                section.right_margin  = Inches(1.0)

            candidate_name = sub.get('candidate_name') or 'Unknown'
            total_marks    = sub.get('total_marks_obtained')
            percentage     = sub.get('percentage')
            submitted_at   = sub.get('submitted_at', '')
            raw_decision   = (sub.get('candidate_decision') or 'pending').lower()
            decision_label = raw_decision.capitalize()
            total_possible = sum(q.get('marks', 0) for q in questions.values())

            # ── Title ────────────────────────────────────────────────────────
            title_para = doc.add_paragraph()
            title_run  = title_para.add_run(title)
            title_run.bold       = True
            title_run.font.size  = Pt(18)
            title_run.font.color.rgb = RGBColor(0x1e, 0x3a, 0x8a)  # indigo-900
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # ── Candidate header ─────────────────────────────────────────────
            name_para = doc.add_paragraph()
            name_run  = name_para.add_run(candidate_name)
            name_run.bold      = True
            name_run.font.size = Pt(15)
            name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()  # spacer

            # ── Summary table ────────────────────────────────────────────────
            tbl = doc.add_table(rows=1, cols=2)
            tbl.style = 'Table Grid'
            tbl.columns[0].width = Inches(2.2)
            tbl.columns[1].width = Inches(4.0)

            def _tbl_row(label, value, color=None):
                row  = tbl.add_row()
                lbl  = row.cells[0].paragraphs[0]
                lbl_run = lbl.add_run(label)
                lbl_run.bold = True
                lbl_run.font.size = Pt(10)
                val  = row.cells[1].paragraphs[0]
                val_run = val.add_run(value)
                val_run.font.size = Pt(10)
                if color:
                    val_run.font.color.rgb = color

            # remove blank first row
            tbl.rows[0]._element.getparent().remove(tbl.rows[0]._element)

            _tbl_row("Email",     sub.get('candidate_email', ''))
            if total_marks is not None:
                score_str = f"{percentage:.1f}%  ({total_marks} / {total_possible} marks)"
            else:
                score_str = "Not evaluated"
            _tbl_row("Score",     score_str)
            _tbl_row("Submitted", submitted_at)
            _tbl_row("Decision",  decision_label,
                     DECISION_COLORS.get(raw_decision, RGBColor(0x6b, 0x72, 0x80)))

            doc.add_paragraph()  # spacer

            # ── Answers ──────────────────────────────────────────────────────
            for idx, ans in enumerate(answers, 1):
                q_id       = ans.get('question_id', '')
                q          = questions.get(q_id, {})
                q_text     = q.get('question_text', 'Unknown question')
                difficulty = q.get('difficulty', '').capitalize()
                q_marks    = q.get('marks', 0)
                awarded    = ans.get('marks_awarded')
                lang       = ans.get('programming_language', '')
                code       = (ans.get('submitted_code') or '').strip()
                feedback   = (ans.get('feedback') or '').strip()
                covered    = ans.get('key_points_covered') or []
                missed     = ans.get('key_points_missed') or []
                quality    = ans.get('code_quality_score')
                awarded_str = str(awarded) if awarded is not None else 'N/A'

                # Question heading
                q_para = doc.add_paragraph()
                q_run  = q_para.add_run(f"Q{idx}. {q_text}")
                q_run.bold      = True
                q_run.font.size = Pt(12)
                q_run.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)  # blue-800

                # Meta line
                meta_para = doc.add_paragraph()
                meta_para.add_run(f"Difficulty: {difficulty}   |   Marks: {awarded_str} / {q_marks}   |   Language: {lang}").font.size = Pt(10)

                # Code block
                if code:
                    code_label = doc.add_paragraph()
                    code_label.add_run("Submitted Code").bold = True

                    code_para = doc.add_paragraph(code)
                    code_para.style = doc.styles['No Spacing']
                    for run in code_para.runs:
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
                    # Light gray shading
                    from docx.oxml.ns import qn
                    from docx.oxml   import OxmlElement
                    pPr = code_para._p.get_or_add_pPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'),   'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'),  'F3F4F6')
                    pPr.append(shd)

                # AI Feedback
                if feedback:
                    fb_label = doc.add_paragraph()
                    fb_label.add_run("AI Feedback").bold = True
                    doc.add_paragraph(feedback)

                if covered:
                    p = doc.add_paragraph()
                    _add_label_value(p, "Key Points Covered", ", ".join(covered))
                if missed:
                    p = doc.add_paragraph()
                    _add_label_value(p, "Key Points Missed", ", ".join(missed))
                if quality is not None:
                    p = doc.add_paragraph()
                    _add_label_value(p, "Code Quality Score", f"{quality}/100")

                # Divider
                doc.add_paragraph("─" * 60)

            buf = BytesIO()
            doc.save(buf)
            return buf.getvalue()

        # Build ZIP in memory
        zip_buffer = BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
            for sub in submissions:
                sub_id = sub['id']
                candidate_name = sub.get('candidate_name') or 'Unknown'
                safe_candidate = _sanitize_name(candidate_name)
                
                # Check for duplicate names and append email if needed
                if name_counts.get(safe_candidate.lower(), 0) > 1:
                    email = sub.get('candidate_email')
                    if email:
                        safe_email = re.sub(r'[^a-zA-Z0-9_\-\.]', '_', email)
                        safe_candidate = f"{safe_candidate}_{safe_email}"
                    else:
                        safe_candidate = f"{safe_candidate}_{sub_id[-6:]}"

                folder = f"{safe_title}/{safe_candidate}"

                # Fetch answers for this submission
                answers_result = client.table('coding_answers').select('*').eq(
                    'submission_id', sub_id
                ).execute()
                
                answers_data = answers_result.data or []
                
                # Sort answers by question_number using the questions map for consistency
                answers = sorted(
                    answers_data, 
                    key=lambda a: questions.get(a['question_id'], {}).get('question_number', 999)
                )

                # Build styled Word document
                docx_bytes = _build_docx(sub, answers)
                zf.writestr(f"{folder}/{safe_candidate}_answers.docx", docx_bytes)

                # Download and add resume if present
                resume_path = sub.get('resume_path')
                if resume_path:
                    try:
                        resume_bytes = client.storage.from_('documents').download(resume_path)
                        ext = resume_path.rsplit('.', 1)[-1] if '.' in resume_path else 'pdf'
                        zf.writestr(f"{folder}/{safe_candidate}_resume.{ext}", resume_bytes)
                    except Exception as resume_err:
                        logger.warning(f"Could not download resume for {sub_id}: {resume_err}")

        zip_buffer.seek(0)
        safe_filename = re.sub(r'[^\w\-]', '_', title)

        return StreamingResponse(
            zip_buffer,
            media_type='application/zip',
            headers={
                'Content-Disposition': f'attachment; filename="{safe_filename}_submissions.zip"'
            }
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error exporting submissions: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to export submissions: {str(e)}"
        )


# ============================================================================
# New endpoints: Edit, Clone, Send Invites, Bulk Decision, Bulk Delete, Notes
# ============================================================================

@router.patch("/{interview_id}", summary="Edit interview details")
async def update_interview(
    interview_id: str,
    body: InterviewUpdate,
    ctx: OrgContext = Depends(require_permission('interview:create'))
):
    """Update interview title, description, scheduled times, or grace period."""
    try:
        service = get_coding_interview_service()
        result = await service.update_interview(
            interview_id=interview_id,
            update_data=body.dict(exclude_none=True),
            user_id=ctx.user_id,
        )
        return result
    except ValueError as e:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=str(e))
    except Exception as e:
        logger.error(f"Error updating interview {interview_id}: {e}")
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="Failed to update interview")


@router.post("/{interview_id}/clone", summary="Clone interview")
async def clone_interview(
    interview_id: str,
    ctx: OrgContext = Depends(require_permission('interview:create'))
):
    """Create a copy of an interview with all its questions and a new access token."""
    try:
        service = get_coding_interview_service()
        result = await service.clone_interview(interview_id, ctx.user_id, ctx.org_id)
        return result
    except ValueError as e:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=str(e))
    except Exception as e:
        logger.error(f"Error cloning interview {interview_id}: {e}")
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="Failed to clone interview")


@router.post("/{interview_id}/send-invites", summary="Send invite emails to candidates")
async def send_invites(
    interview_id: str,
    ctx: OrgContext = Depends(require_permission('interview:create'))
):
    """Email the interview link to all pre-registered candidates who haven't submitted."""
    try:
        service = get_coding_interview_service()
        result = await service.send_interview_invites(interview_id, ctx.user_id)
        return result
    except ValueError as e:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=str(e))
    except Exception as e:
        logger.error(f"Error sending invites for interview {interview_id}: {e}")
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="Failed to send invites")


@router.post("/{interview_id}/bulk-decision", summary="Bulk set decision on submissions")
async def bulk_decision(
    interview_id: str,
    body: BulkDecisionRequest,
    ctx: OrgContext = Depends(require_permission('interview:evaluate'))
):
    """Set the same decision (advanced/rejected/hold/pending) on multiple submissions at once."""
    try:
        service = get_coding_interview_service()
        result = await service.bulk_submission_decision(
            interview_id=interview_id,
            submission_ids=body.submission_ids,
            decision=body.decision,
            decided_by=ctx.user_id,
        )
        return result
    except ValueError as e:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=str(e))
    except Exception as e:
        logger.error(f"Error bulk decision for interview {interview_id}: {e}")
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="Failed to update decisions")


@router.post("/{interview_id}/candidates/bulk-delete", summary="Bulk delete pre-registered candidates")
async def bulk_delete_candidates(
    interview_id: str,
    body: BulkDeleteCandidatesRequest,
    ctx: OrgContext = Depends(require_permission('interview:create'))
):
    """Delete multiple pre-registered candidates from the interview pipeline."""
    try:
        service = get_coding_interview_service()
        result = await service.bulk_delete_candidates(
            interview_id=interview_id,
            candidate_ids=body.candidate_ids,
            user_id=ctx.user_id,
        )
        return result
    except ValueError as e:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=str(e))
    except Exception as e:
        logger.error(f"Error bulk deleting candidates for interview {interview_id}: {e}")
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="Failed to delete candidates")


@router.patch(
    "/submissions/{submission_id}/answers/{answer_id}/notes",
    summary="Save evaluator notes and optional score override",
)
async def save_evaluator_notes(
    submission_id: str,
    answer_id: str,
    body: EvaluatorNotesUpdate,
    ctx: OrgContext = Depends(require_permission('interview:evaluate'))
):
    """Save interviewer notes and optionally override the AI-assigned marks for a specific answer."""
    try:
        service = get_coding_interview_service()
        result = await service.save_evaluator_notes(
            submission_id=submission_id,
            answer_id=answer_id,
            notes=body.notes,
            marks_override=body.marks_override,
            evaluator_id=ctx.user_id,
        )
        return result
    except ValueError as e:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=str(e))
    except Exception as e:
        logger.error(f"Error saving evaluator notes for answer {answer_id}: {e}")
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="Failed to save notes")


@router.delete("/submissions/{submission_id}", summary="Delete a submission")
async def delete_submission(
    submission_id: str,
    ctx: OrgContext = Depends(require_permission('interview:create'))
):
    """Delete a candidate submission (any status: in_progress, submitted, etc.). Requires interview ownership."""
    try:
        service = get_coding_interview_service()
        await service.delete_submission(
            submission_id=submission_id,
            user_id=ctx.user_id
        )
        return {'message': 'Submission deleted successfully'}
    except ValueError as e:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=str(e))
    except Exception as e:
        logger.error(f"Error deleting submission {submission_id}: {e}")
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=f"Failed to delete submission: {str(e)}")
